#!/usr/bin/env python3
import errno
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from bs4 import BeautifulSoup
from PIL import Image
from docx.oxml.shared import OxmlElement, qn
from html.parser import HTMLParser
from PreprocessMarkdown2docx import PreprocessMarkdown2docx
'''
    For docx - see
    https://python-docx.readthedocs.io/en/latest/index.html
    For styles - see 
    https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
    https://python-docx.readthedocs.io/en/latest/user/styles-using.html
    https://bettersolutions.com/word/styles/list-all-built-in-styles.htm
    For Pictures and more - see
    https://python-docx.readthedocs.io/en/latest/user/quickstart.html
    For Table of Contents:
    https://github.com/python-openxml/python-docx/issues/36
    For Linux error codes:
    https://www.thegeekstuff.com/2010/10/linux-error-codes/
    For markdown2 ( I abandoned markdown in favour of markdown2 )
    https://www.programcreek.com/python/example/93863/markdown2.markdown
    For bullets:
    https://www.geeksforgeeks.org/working-with-lists-python-docx-module/
    
    This is how you can print all the available styles:
        styles = doc.styles
        for s in styles:
            print(s.name)
'''

purpose = '''
Read valid markdown, and write a nice docx document using basic elements:
* Headings
* Paragraphs
* Pre-formatted text
* Bullets (With one level deep only)
* Tables (No nesting)
* Pictures
'''


def _read_in_markdown(file_name, encoding='utf8'):
    try:
        with open(file_name, 'r', encoding=encoding) as input_fd:
            return input_fd.read()
    except FileNotFoundError as e:
        print(f'ERROR: {e}')
        exit(errno.ENOENT)
    except PermissionError as e:
        print(f'ERROR: {e}')
        exit(errno.EACCES)
    except IsADirectoryError as e:
        print(f'ERROR: {e}')
        exit(errno.EISDIR)


def write_out_html(file_name, text_html, encoding='utf8'):
    with open(file_name, 'w', encoding=encoding) as output_fd:
        output_fd.write(text_html)


def find_page_width(doc):
    return float(doc.sections[0].page_width / 914400)


def do_table_of_contents(document):
    """https://github.com/python-openxml/python-docx/issues/36
    Insert Table of Contents Field. The user will be asked to
    update the field when the docx is opened.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')  # creates a new element
    fld_char.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instr_text.text = r'TOC \o "1-3" \h \z \u'   # change 1-3 depending on heading levels you need
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:t')
    fld_char3.text = "Right-click to update field."
    fld_char2.append(fld_char3)
    fld_char4 = OxmlElement('w:fldChar')
    fld_char4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)


def do_table(doc, table_in, style):
    """Draw a table."""
    the_header = table_in.find('thead')
    the_column_names = the_header.find_all('th')
    the_data = table_in.find_all('td')
    n_cols = len(the_column_names)
    n_rows = int(len(the_data) / n_cols)
    this_table = doc.add_table(rows=n_rows + 1, cols=n_cols, style=style)
    row = this_table.rows[0].cells
    for h_index, header in enumerate(the_column_names):
        row[h_index].text = '' if header.text == '' else header.string
    row_index = 0
    for d_index, data in enumerate(the_data):
        if not d_index % n_cols:
            row = this_table.rows[row_index + 1].cells
            row_index += 1
        row[d_index % n_cols].text = '' if data.text == '' else data.string


def find_image_size(image_file):
    return Image.open(image_file).size


def do_paragraph(line,
                 doc,
                 page_width_inches,
                 style_body,
                 assumed_pixels_per_inch=200,
                 picture_fraction_of_width=0.7):
    is_image = line.find('img')
    if is_image is not None:
        image_source = is_image['src']
        w, h = find_image_size(image_source)
        w_in_inches = w / assumed_pixels_per_inch
        picture_width_inches = page_width_inches * picture_fraction_of_width
        chosen_width = min(picture_width_inches, w_in_inches)
        doc.add_picture(image_source,
                        width=docx.shared.Inches(chosen_width))
        return
    paragraph = doc.add_paragraph(line.text.strip(), style=style_body)
    paragraph.style.font.name = 'Calibri'
    paragraph.style.font.size = Pt(11)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)


def do_pre_code(line, doc, style_quote_table):
    table = doc.add_table(rows=1, cols=1, style=style_quote_table)
    cell = table.cell(0, 0)
    cell.text = line.text
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0]
    run_obj = paragraph.runs
    run = run_obj[0]
    font = run.font
    font.size = Pt(10)
    font.name = 'Verdana'


def do_fake_horizontal_rule(doc, length_of_line=80, c='_'):
    paragraph = doc.add_paragraph(c * length_of_line)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _eat_soup(soup,
              doc,
              page_width_inches,
              style_quote_table,
              style_body,
              style_table,
              heading_style,
              table_of_contents_string='contents'):
    """HTML from markdown has been converted to a beautiful soup (bs4) object.
    Process the object to render a Word docx."""
    list_of_tables = soup.find_all('table')
    table_counter = 0
    table_of_contents_done = 0
    for line in soup:
        if line is not None:
            if str(line).lower().find(table_of_contents_string) >= 0 and table_of_contents_done < 2:
                table_of_contents_done += 1
                if table_of_contents_done == 2:
                    do_table_of_contents(doc)
            if line.find('em'):
                try:
                    doc.add_paragraph(line.text.strip(), style=style_body)
                except AttributeError:
                    pass
                continue
            if str(line).lower().find('<hr/>') == 0:
                do_fake_horizontal_rule(doc)
                continue
            if line.name == 'h1':
                heading = doc.add_heading(line, 0)
                heading.style = heading_style
                continue
            if line.name == 'h2':
                heading = doc.add_heading(line, 1)
                heading.style = heading_style
                continue
            if line.name == 'h3':
                heading = doc.add_heading(line, 2)
                heading.style = heading_style
                continue
            if line.name == 'h4':
                heading = doc.add_heading(line, 3)
                heading.style = heading_style
                continue
            if line.name == 'p':
                paragraph = do_paragraph(line, doc, page_width_inches, style_body)
                continue
            if line.name == 'pre':
                do_pre_code(line, doc, style_quote_table)
            if line.name == 'table':
                this_table_in = list_of_tables[table_counter]
                do_table(doc, this_table_in, style_table)
                table_counter += 1
                continue
            if line.name == 'ul':  # unordered list (bullets)
                parser = HtmlListParser()
                parser.doc = doc
                parser.feed(str(line))
                continue
            if line.name == 'ol':  # ordered list (numbered)
                parser = HtmlListParser()
                HtmlListParser.lists = ['List Number', 'List Number 2', 'List Number 3']
                HtmlListParser.spare_list = '#  '
                parser.doc = doc
                parser.feed(str(line))
                continue
    return doc


class HtmlListParser(HTMLParser):
    list_level = -1
    lists = ['List Bullet', 'List Bullet 2', 'List Bullet 3']
    doc = None  # the .docx document object
    spacing = '    '  # used if we run out of bullet levels
    spare_list = '○  '

    def handle_starttag(self, tag, attrs):
        if tag in ['ol', 'ul']:
            self.list_level += 1

    def handle_endtag(self, tag):
        if tag in ['ol', 'ul']:
            self.list_level -= 1

    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.list_level in range(len(self.lists)):
                self.doc.add_paragraph(data, self.lists[self.list_level])
            else:
                self.doc.add_paragraph('        ' + self.spacing * self.list_level + self.spare_list + data)


class Markdown2docx:
    style_table = 'Medium Shading 1 Accent 3'
    style_quote = 'Body Text'
    style_body = 'Body Text'
    style_quote_table = 'Table Grid'
    toc_indicator = 'contents'



    def __init__(self, project, markdown=None, file_stream=None):
        self.infile = '.'.join([project, 'md'])
        self.outfile = '.'.join([project, 'docx'])
        self.html_out_file = '.'.join([project, 'html'])
        self.project = project
        self.doc = docx.Document()

        self.heading_style = self.doc.styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
        self.heading_style.base_style = self.doc.styles['Heading 2']
        self.heading_style.font.color.rgb = RGBColor(0, 0, 0)
        self.heading_style.font.size = Pt(18)
        self.heading_style.font.name = 'Verdana'
        self.heading_style.paragraph_format.space_before = Pt(12)
        self.heading_style.paragraph_format.space_after = Pt(6)


        self.file_stream = file_stream
        self.page_width_inches = find_page_width(self.doc)
        # self.html = markdown.markdown(_read_in_markdown(self.infile), extensions=['tables'])
        self.markdown = markdown
        self.html = markdown2.markdown(self.markdown, extras=[
            'fenced-code-blocks',
            'code-friendly',
            'wiki-tables',
            'tables'
        ])

        self.soup = BeautifulSoup(self.html, 'html.parser')
        # return self.soup

    def eat_soup(self):
        _eat_soup(self.soup,
                  self.doc,
                  self.page_width_inches,
                  self.style_quote_table,
                  self.style_body,
                  self.style_table,
                  self.heading_style,
                  table_of_contents_string=self.toc_indicator)

    def __del__(self):
        pass

    def __repr__(self):
        return f'Markdown2docx("{self.project}")'

    def __str__(self):
        return "'{" \
               f"project':{self.project}," \
               f"'style_table':{self.style_table}," \
               f"'style_quote':{self.style_quote}," \
               f"'style_body':{self.style_body}," \
               f"'style_quote_table':{self.style_quote_table}," \
               f"'toc_indicator':{self.toc_indicator}" \
               "}"

    def styles(self):
        return {
               'project': {self.project},
               'style_table': {self.style_table},
               'style_quote': {self.style_quote},
               'style_body': {self.style_body},
               'style_quote_table': {self.style_quote_table},
               'toc_indicator': {self.toc_indicator}
               }

    def write_html(self):
        write_out_html(self.html_out_file, self.html)

    def save(self):
        self.doc.save(self.file_stream)


def __main__(project):
    ppm2w = PreprocessMarkdown2docx(project)
    markdown = ppm2w.get_all_but_macros()
    markdown = ppm2w.do_substitute_tokens(markdown)
    markdown = ppm2w.do_execute_commands(markdown)
    # print(markdown)
    # for i in markdown:
    #    print(i)
    # for k, v in macros.items():
    #    print(k,v)
    # project = Markdown2docx('hello')
    project = Markdown2docx(project, markdown)
    project.eat_soup()
    # project.write_html()  # optional
    # print(type(project.styles()))
    # for k, v in project.styles().items():
    #     print(f'stylename: {k} = {v}')
    project.save()


if __name__ == "__main__":
    __main__('hello')
